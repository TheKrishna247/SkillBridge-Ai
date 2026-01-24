import pdfplumber
import io
import docx
import pytesseract
from PIL import Image
import re

def extract_text_from_pdf(file):
    """Extract text from PDF file"""
    text = ""
    try:
        with pdfplumber.open(file) as pdf:
            for page in pdf.pages:
                text += page.extract_text() + "\n"
    except:
        # Try alternative method
        import PyPDF2
        reader = PyPDF2.PdfReader(file)
        for page in reader.pages:
            text += page.extract_text() + "\n"
    return text.strip()

def extract_text_from_docx(file):
    doc = docx.Document(file)
    text_parts = []

    # Paragraph text
    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text)

    # Table text (CRITICAL FIX)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text_parts.append(cell.text)

    return "\n".join(text_parts).strip()


def extract_text_from_image(file):
    """Extract text from image using OCR"""
    image = Image.open(file)
    text = pytesseract.image_to_string(image)
    return text.strip()

def extract_text_from_txt(file):
    """Extract text from text file"""
    return file.read().decode('utf-8')

def extract_text(file):
    """
    Main function to extract text from any file format
    Streamlit-safe: reads file ONCE and routes by extension
    """
    file.seek(0)
    file_bytes = file.read()
    file_type = file.name.lower()

    if file_type.endswith('.pdf'):
        return extract_text_from_pdf(io.BytesIO(file_bytes))

    elif file_type.endswith('.docx'):
        return extract_text_from_docx(io.BytesIO(file_bytes))

    elif file_type.endswith(('.png', '.jpg', '.jpeg', '.bmp', '.tiff')):
        return extract_text_from_image(io.BytesIO(file_bytes))

    elif file_type.endswith('.txt'):
        return file_bytes.decode("utf-8", errors="ignore")

    else:
        return file_bytes.decode("utf-8", errors="ignore")


def extract_skills_from_text(text):
    """Extract skills from resume text"""
    # Common skills database
    common_skills = [
        "python", "javascript", "java", "html", "css", "react", "node", "sql",
        "mongodb", "aws", "docker", "git", "machine learning", "data analysis",
        "excel", "tableau", "power bi", "linux", "networking", "cybersecurity"
    ]
    
    text_lower = text.lower()
    found_skills = []
    
    for skill in common_skills:
        if skill in text_lower:
            found_skills.append(skill.title())
    
    return found_skills